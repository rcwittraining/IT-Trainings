#!/usr/bin/env python3
# RCW-NixTrace v1.0.0 — Linux Server Inventory & Documentation Tool
# Windows portable GUI (Tkinter) that SSHes to a Linux server, gathers a
# complete inventory, and exports it to a Word (.docx) document.
import tkinter as tk
from tkinter import ttk, scrolledtext, filedialog, messagebox
import paramiko
import os, sys, datetime, re, base64

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except Exception:
    DOCX_OK = False

APP_NAME = "RCW-NixTrace"
APP_VER = "1.0.0"
APP_TAG = "Linux Server Inventory & Documentation Tool"

# Maximum characters kept per command output (keeps the document usable).
MAX_OUT = 8000

# ---------------------------------------------------------------------------
# Command catalogue (section -> list of (command, needs_root))
# ---------------------------------------------------------------------------
SECTIONS = [
    ("System Overview", [
        ("uname -a", False),
        ("cat /etc/os-release", False),
        ("hostnamectl 2>/dev/null", False),
        ("date", False),
        ("uptime", False),
        ("cat /proc/uptime", False),
        ("cat /proc/loadavg", False),
        ("last reboot | head -5", False),
        ("who", False),
        ("w 2>/dev/null", False),
    ]),
    ("CPU", [
        ("lscpu", False),
        ("nproc", False),
        ("cat /proc/cpuinfo | head -40", False),
    ]),
    ("Memory", [
        ("free -h", False),
        ("cat /proc/meminfo", False),
    ]),
    ("Storage & Disks", [
        ("lsblk -o NAME,SIZE,TYPE,MOUNTPOINT,FSTYPE,LABEL,MODEL", False),
        ("df -h", False),
        ("cat /etc/fstab", False),
        ("findmnt -D 2>/dev/null", False),
        ("ls -la /dev/disk/by-label 2>/dev/null", False),
    ]),
    ("Network", [
        ("ip addr", False),
        ("ip route", False),
        ("ip -6 route", False),
        ("cat /etc/resolv.conf", False),
        ("cat /etc/hosts", False),
        ("ss -tulpn 2>/dev/null", True),
        ("netstat -tulpn 2>/dev/null", False),
    ]),
    ("Firewall & Security", [
        ("iptables -L -n -v 2>/dev/null", True),
        ("nft list ruleset 2>/dev/null", True),
        ("firewall-cmd --list-all 2>/dev/null", True),
        ("ufw status verbose 2>/dev/null", True),
        ("sestatus 2>/dev/null", False),
        ("getenforce 2>/dev/null", False),
        ("cat /etc/ssh/sshd_config 2>/dev/null", False),
    ]),
    ("Users & Groups", [
        ("cat /etc/passwd", False),
        ("cat /etc/group", False),
        ("getent passwd", False),
        ("lastlog 2>/dev/null | head -60", False),
        ("ls -la /etc/sudoers.d 2>/dev/null", False),
        ("cat /etc/sudoers 2>/dev/null", False),
    ]),
    ("Services (systemd)", [
        ("systemctl list-units --type=service --state=running --no-pager 2>/dev/null", False),
        ("systemctl list-units --type=service --state=failed --no-pager 2>/dev/null", False),
        ("systemctl list-unit-files --type=service --state=enabled --no-pager 2>/dev/null", False),
    ]),
    ("Installed Packages", [
        ("echo RPM_COUNT=$(rpm -qa 2>/dev/null | wc -l)", False),
        ("echo DPKG_COUNT=$(dpkg -l 2>/dev/null | wc -l)", False),
        ("rpm -qa 2>/dev/null", False),
        ("dpkg-query -W -f='${Package} ${Version}\\n' 2>/dev/null", False),
    ]),
    ("Hardware", [
        ("lspci 2>/dev/null", False),
        ("lsusb 2>/dev/null", False),
        ("dmidecode -t system 2>/dev/null", True),
        ("dmidecode -t memory 2>/dev/null", True),
        ("lsblk -d -o NAME,SIZE,MODEL,ROTA 2>/dev/null", False),
    ]),
    ("Scheduling & Jobs", [
        ("ls -la /etc/cron* 2>/dev/null", False),
        ("ls -la /var/spool/cron 2>/dev/null", False),
        ("cat /etc/crontab 2>/dev/null", False),
    ]),
    ("Kernel & Modules", [
        ("sysctl -a 2>/dev/null", False),
        ("lsmod", False),
    ]),
    ("Environment", [
        ("printenv", False),
    ]),
]


# ---------------------------------------------------------------------------
# SSH probe
# ---------------------------------------------------------------------------
class SSHProbe:
    def __init__(self):
        self.client = None
        self.ip = None

    def connect(self, host, port, user, auth, secret, timeout=15):
        self.client = paramiko.SSHClient()
        self.client.set_missing_host_key_policy(paramiko.AutoAddPolicy())
        if auth == "key":
            pkey = None
            try:
                if secret.endswith(".ppk"):
                    pkey = paramiko.RSAKey.from_private_key_file(secret)
                else:
                    pkey = paramiko.Ed25519Key.from_private_key_file(secret)
            except Exception:
                try:
                    pkey = paramiko.RSAKey.from_private_key_file(secret)
                except Exception:
                    pkey = paramiko.ECDSAKey.from_private_key_file(secret)
            self.client.connect(hostname=host, port=port, username=user,
                                pkey=pkey, timeout=timeout, look_for_keys=False,
                                allow_agent=False)
        else:
            self.client.connect(hostname=host, port=port, username=user,
                                password=secret, timeout=timeout,
                                look_for_keys=False, allow_agent=False)
        self.ip = self.client.get_transport().getpeername()[0]
        return True

    def cmd(self, command, sudo_pw=None, timeout=120):
        if not self.client:
            return ("", "not connected", -1)
        try:
            if sudo_pw:
                full = 'sudo -S -p "" %s' % command
                stdin, stdout, stderr = self.client.exec_command(full, timeout=timeout)
                try:
                    stdin.write(sudo_pw + "\n")
                    stdin.flush()
                except Exception:
                    pass
            else:
                stdin, stdout, stderr = self.client.exec_command(command, timeout=timeout)
            out = stdout.read().decode("utf-8", "replace")
            err = stderr.read().decode("utf-8", "replace")
            try:
                code = stdout.channel.recv_exit_status()
            except Exception:
                code = 0
            return (out, err, code)
        except Exception as e:
            return ("", str(e), -1)

    def close(self):
        try:
            self.client.close()
        except Exception:
            pass


# ---------------------------------------------------------------------------
# Small parsers for the summary table
# ---------------------------------------------------------------------------
def parse_os_release(text):
    name = rel = ""
    for line in text.splitlines():
        if line.startswith("PRETTY_NAME="):
            name = line.split("=", 1)[1].strip().strip('"')
        if line.startswith("VERSION="):
            rel = line.split("=", 1)[1].strip().strip('"')
    return (name + (" " + rel if rel else "")).strip() or "Unknown"


def parse_cpu(text):
    model = sockets = cores = ""
    for line in text.splitlines():
        if line.startswith("Model name:"):
            model = line.split(":", 1)[1].strip()
        if line.startswith("Socket(s):"):
            sockets = line.split(":", 1)[1].strip()
        if line.startswith("CPU(s):"):
            cores = line.split(":", 1)[1].strip()
    return "%s (%s vCPU, %s socket(s))" % (model or "CPU", cores or "?", sockets or "?")


def parse_mem(text):
    for line in text.splitlines():
        if line.startswith("MemTotal:"):
            kb = int(line.split()[1])
            return "%.1f GB" % (kb / 1024.0 / 1024.0)
    return "Unknown"


def gather_meta(probe):
    meta = {}
    uname = probe.cmd("uname -a")[0]
    parts = uname.split()
    meta["kernel"] = parts[2] if len(parts) > 2 else ""
    meta["arch"] = probe.cmd("uname -m")[0].strip()
    meta["os"] = parse_os_release(probe.cmd("cat /etc/os-release")[0])
    meta["hostname"] = probe.cmd("hostname")[0].strip() or (parts[1] if len(parts) > 1 else "unknown")
    up = probe.cmd("uptime -p 2>/dev/null")[0].strip()
    if not up:
        up = probe.cmd("cat /proc/uptime")[0].strip()
    meta["uptime"] = up
    meta["cpu"] = parse_cpu(probe.cmd("lscpu")[0])
    meta["mem"] = parse_mem(probe.cmd("cat /proc/meminfo")[0])
    meta["date"] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    meta["ip"] = probe.ip or ""
    meta["host"] = probe._host if hasattr(probe, "_host") else meta["hostname"]
    return meta


# ---------------------------------------------------------------------------
# Collection
# ---------------------------------------------------------------------------
def collect(probe, sudo_pw, log_fn):
    results = {}
    total = sum(len(cmds) for _, cmds in SECTIONS)
    done = 0
    for sec, cmds in SECTIONS:
        results[sec] = []
        for cmd, needs_root in cmds:
            out, err, code = probe.cmd(cmd, sudo_pw if needs_root else None)
            results[sec].append((cmd, out, err, code))
            done += 1
            log_fn("[%d/%d] %s -> %d bytes" % (done, total, cmd, len(out)))
    return results


# ---------------------------------------------------------------------------
# Word document
# ---------------------------------------------------------------------------
def build_document(results, meta, path):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    doc.add_heading("Linux Server Inventory Report", level=0)
    sub = doc.add_paragraph()
    r = sub.add_run("%s v%s  -  %s" % (APP_NAME, APP_VER, APP_TAG))
    r.italic = True
    r.font.size = Pt(11)

    doc.add_heading("Scan Summary", level=1)
    rows = [
        ("Target Host", meta.get("host")),
        ("Resolved IP", meta.get("ip")),
        ("Scan Date", meta.get("date")),
        ("Operating System", meta.get("os")),
        ("Kernel", meta.get("kernel")),
        ("Architecture", meta.get("arch")),
        ("Uptime", meta.get("uptime")),
        ("CPU", meta.get("cpu")),
        ("Memory (total)", meta.get("mem")),
        ("Generated By", "%s v%s" % (APP_NAME, APP_VER)),
    ]
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light Grid Accent 1"
    for k, v in rows:
        c = tbl.add_row().cells
        c[0].text = str(k)
        c[1].text = str(v if v is not None else "")

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("This document was generated automatically by %s. " % APP_NAME)
    note.add_run("Sections below contain the raw output of standard Linux "
                 "inventory commands captured over SSH.").italic = True

    for sec, items in results.items():
        doc.add_heading(sec, level=1)
        for cmd, out, err, code in items:
            doc.add_heading(cmd, level=2)
            text = out.strip()
            if not text:
                text = (err.strip() or "(no output / command unavailable or permission denied)")
            if len(text) > MAX_OUT:
                text = text[:MAX_OUT] + "\n... [output truncated; %d characters omitted]" % (len(text) - MAX_OUT)
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.font.name = "Consolas"
            run.font.size = Pt(8)

    doc.save(path)


# ---------------------------------------------------------------------------
# GUI
# ---------------------------------------------------------------------------
class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("%s v%s" % (APP_NAME, APP_VER))
        try:
            self.iconbitmap(os.path.join(os.path.dirname(os.path.abspath(__file__)), "rcw.ico"))
        except Exception:
            pass
        self.probe = None
        self.results = None
        self.meta = None
        self.sudo_pw = ""
        self.build_ui()

    def build_ui(self):
        pad = {"padx": 8, "pady": 4}

        conn = ttk.LabelFrame(self, text="Target Linux Server (SSH)")
        conn.pack(fill="x", **pad)

        ttk.Label(conn, text="Host:").grid(row=0, column=0, sticky="e")
        self.host = ttk.Entry(conn, width=26); self.host.grid(row=0, column=1, **pad)
        self.host.insert(0, "192.168.1.10")
        ttk.Label(conn, text="Port:").grid(row=0, column=2, sticky="e")
        self.port = ttk.Entry(conn, width=8); self.port.grid(row=0, column=3, **pad)
        self.port.insert(0, "22")

        ttk.Label(conn, text="Username:").grid(row=1, column=0, sticky="e")
        self.user = ttk.Entry(conn, width=26); self.user.grid(row=1, column=1, **pad)
        self.user.insert(0, "root")

        ttk.Label(conn, text="Auth:").grid(row=1, column=2, sticky="e")
        self.auth = tk.StringVar(value="password")
        ttk.Radiobutton(conn, text="Password", variable=self.auth, value="password",
                        command=self.on_auth).grid(row=1, column=3, sticky="w")
        ttk.Radiobutton(conn, text="Key file", variable=self.auth, value="key",
                        command=self.on_auth).grid(row=1, column=4, sticky="w")

        ttk.Label(conn, text="Secret:").grid(row=2, column=0, sticky="e")
        self.secret = ttk.Entry(conn, width=40, show="*"); self.secret.grid(row=2, column=1, columnspan=3, **pad)
        self.browse = ttk.Button(conn, text="Browse...", command=self.on_browse)
        self.browse.grid(row=2, column=4, **pad)

        ttk.Label(conn, text="Sudo pw:").grid(row=3, column=0, sticky="e")
        self.sudo = ttk.Entry(conn, width=40, show="*"); self.sudo.grid(row=3, column=1, columnspan=3, **pad)
        ttk.Label(conn, text="(optional)").grid(row=3, column=4, sticky="w")

        self.connect_btn = ttk.Button(conn, text="Connect", command=self.on_connect)
        self.connect_btn.grid(row=4, column=1, **pad)
        self.disconnect_btn = ttk.Button(conn, text="Disconnect", command=self.on_disconnect, state="disabled")
        self.disconnect_btn.grid(row=4, column=2, **pad)
        self.status = ttk.Label(conn, text="Not connected"); self.status.grid(row=4, column=3, columnspan=2, sticky="w")

        actions = ttk.Frame(self)
        actions.pack(fill="x", **pad)
        self.scan_btn = ttk.Button(actions, text="Run Full Inventory", command=self.on_scan, state="disabled")
        self.scan_btn.pack(side="left", **pad)
        self.export_btn = ttk.Button(actions, text="Export Word (.docx)", command=self.on_export, state="disabled")
        self.export_btn.pack(side="left", **pad)
        self.progress = ttk.Label(actions, text="")
        self.progress.pack(side="left", **pad)

        self.log = scrolledtext.ScrolledText(self, height=18, width=110, font=("Consolas", 9))
        self.log.pack(fill="both", expand=True, **pad)

        self.log.insert("end", "%s v%s ready.\n" % (APP_NAME, APP_VER))
        if not DOCX_OK:
            self.log.insert("end", "WARNING: python-docx not available - export disabled.\n")

    def on_auth(self):
        self.secret.config(show="" if self.auth.get() == "key" else "*")

    def on_browse(self):
        p = filedialog.askopenfilename(title="Select private key file",
                                       filetypes=[("Key files", "*.pem *.ppk *.key *.rsa *.ed25519"), ("All", "*.*")])
        if p:
            self.secret.delete(0, "end"); self.secret.insert(0, p)

    def on_connect(self):
        self.status.config(text="Connecting...")
        self.connect_btn.config(state="disabled")
        self.log_insert("Connecting to %s:%s as %s ...\n" %
                        (self.host.get(), self.port.get(), self.user.get()))
        threading_dummy(self._do_connect)

    def _do_connect(self):
        try:
            p = SSHProbe()
            p._host = self.host.get()
            p.connect(self.host.get(), int(self.port.get() or 22), self.user.get(),
                      self.auth.get(), self.secret.get())
            self.probe = p
            self.meta = gather_meta(p)
            self.after(0, lambda: self._connected_ok())
        except Exception as e:
            self.after(0, lambda: self._connected_fail(str(e)))

    def _connected_ok(self):
        self.status.config(text="Connected: %s (%s)" % (self.meta.get("hostname"), self.meta.get("ip")))
        self.connect_btn.config(state="disabled")
        self.disconnect_btn.config(state="normal")
        self.scan_btn.config(state="normal")
        self.sudo_pw = self.sudo.get()
        self.log_insert("Connected. OS: %s | Kernel: %s\n" %
                        (self.meta.get("os"), self.meta.get("kernel")))

    def _connected_fail(self, err):
        self.status.config(text="Connection failed")
        self.connect_btn.config(state="normal")
        self.log_insert("ERROR: %s\n" % err)

    def on_disconnect(self):
        if self.probe:
            self.probe.close()
            self.probe = None
        self.status.config(text="Not connected")
        self.connect_btn.config(state="normal")
        self.disconnect_btn.config(state="disabled")
        self.scan_btn.config(state="disabled")
        self.export_btn.config(state="disabled")
        self.log_insert("Disconnected.\n")

    def on_scan(self):
        self.scan_btn.config(state="disabled")
        self.export_btn.config(state="disabled")
        self.log_insert("Starting full inventory scan...\n")
        threading_dummy(self._do_scan)

    def _do_scan(self):
        try:
            self.results = collect(self.probe, self.sudo_pw or None, self.log_insert)
            self.after(0, self._scan_done)
        except Exception as e:
            self.after(0, lambda: self.log_insert("SCAN ERROR: %s\n" % e))

    def _scan_done(self):
        self.scan_btn.config(state="normal")
        self.export_btn.config(state="normal")
        self.log_insert("Inventory complete. %d sections collected.\n" % len(self.results))
        self.log_insert("Click 'Export Word (.docx)' to save the document.\n")

    def on_export(self):
        if not self.results:
            messagebox.showwarning("Nothing to export", "Run the inventory first.")
            return
        if not DOCX_OK:
            messagebox.showerror("Missing dependency", "python-docx is not available.")
            return
        host = (self.meta or {}).get("hostname", "server")
        default = "RCW-NixTrace-%s-%s.docx" % (host, datetime.datetime.now().strftime("%Y%m%d"))
        path = filedialog.asksaveasfilename(
            title="Save inventory document",
            defaultextension=".docx",
            initialfile=default,
            filetypes=[("Word Document", "*.docx"), ("All files", "*.*")])
        if not path:
            return
        try:
            build_document(self.results, self.meta or {}, path)
            self.log_insert("Document saved: %s\n" % path)
            messagebox.showinfo("Saved", "Inventory document saved:\n%s" % path)
            try:
                os.startfile(path)
            except Exception:
                pass
        except Exception as e:
            messagebox.showerror("Export failed", str(e))
            self.log_insert("EXPORT ERROR: %s\n" % e)

    def log_insert(self, msg):
        self.log.insert("end", msg)
        self.log.see("end")


# helper to run a function in a thread without importing threading at top
def threading_dummy(fn):
    import threading
    threading.Thread(target=fn, daemon=True).start()


if __name__ == "__main__":
    App().mainloop()
